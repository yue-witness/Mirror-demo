from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\OneDrive\Mirror")
OUTPUT_DIR = ROOT / "output"
OUTPUT_PATH = OUTPUT_DIR / "PROJECT_MIRROR_Steam正式发售级UI交互改进与最终设计报告.docx"

RUNTIME = ROOT / "demo" / "_qa" / "current-ui-audit" / "runtime"
APPROVED_SLIDES = ROOT / "Documents" / "审核后正式设计" / "DemoUI"
GREEN_SPEC_SLIDES = (
    ROOT
    / "Documents"
    / "AI生成的UI"
    / "PROJECT_MIRROR_UI设计规范_绿色激光版_重设计"
)

INK = RGBColor(0x12, 0x3D, 0x35)
GREEN = RGBColor(0x1D, 0xA9, 0x7B)
GREEN_DARK = RGBColor(0x0B, 0x70, 0x55)
GREEN_PALE = "E8F7F1"
MINT = "F4FBF8"
YELLOW = RGBColor(0xA3, 0x70, 0x00)
YELLOW_PALE = "FFF3CF"
RED = RGBColor(0xA3, 0x2F, 0x42)
RED_PALE = "FDE9ED"
GRAY = RGBColor(0x5A, 0x6B, 0x67)
GRAY_PALE = "EFF3F2"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for side, value in CELL_MARGIN_DXA.items():
        tag = "start" if side == "start" else "end" if side == "end" else side
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_table_borders(table, color: str = "B7D8CC", size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)
        node.set(qn("w:space"), "0")


def set_repeat_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def font_run(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_border_bottom(paragraph, color: str = "1DA97B", size: str = "14") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(paragraph)


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        font_run(lead, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_lead):])
        font_run(rest)
    else:
        run = paragraph.add_run(text)
        font_run(run)


def add_bullet(doc: Document, text: str, level: int = 0, compact: bool = False) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(1.5 if compact else 4)
    paragraph.paragraph_format.line_spacing = 1.08 if compact else 1.25
    run = paragraph.add_run(text)
    font_run(run, size=9.2 if compact else 10.5)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    font_run(run)


def add_callout(
    doc: Document,
    label: str,
    text: str,
    fill: str = GREEN_PALE,
    color: RGBColor = GREEN_DARK,
    trailing_space: bool = True,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color="A9D5C6", size="8")
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    lead = paragraph.add_run(f"{label}  ")
    font_run(lead, size=10.5, bold=True, color=color)
    body = paragraph.add_run(text)
    font_run(body, size=10.5, color=INK)
    if trailing_space:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths_dxa: Sequence[int],
    header_fill: str = GREEN_PALE,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    header_row = table.rows[0]
    mark_header_row(header_row)
    set_repeat_no_split(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        font_run(run, size=9.5, bold=True, color=GREEN_DARK)

    for row_values in rows:
        row = table.add_row()
        set_repeat_no_split(row)
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            font_run(run, size=9.2, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.25) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    inline = run._element.xpath(".//wp:docPr")
    if inline:
        inline[0].set("descr", caption)
        inline[0].set("title", caption)

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_run = caption_paragraph.add_run(caption)
    font_run(caption_run, size=8.5, color=GRAY)


def add_two_figures(doc: Document, left: Path, right: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    left_run = paragraph.add_run()
    left_run.add_picture(str(left), width=Inches(3.12))
    gap = paragraph.add_run("   ")
    font_run(gap)
    right_run = paragraph.add_run()
    right_run.add_picture(str(right), width=Inches(3.12))
    for run, alt in ((left_run, "当前实机界面"), (right_run, "文档设计稿")):
        nodes = run._element.xpath(".//wp:docPr")
        if nodes:
            nodes[0].set("descr", alt)
            nodes[0].set("title", alt)

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    font_run(caption_run, size=8.5, color=GRAY)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0B7055")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    run_properties.append(color)
    run_properties.append(underline)
    run_properties.append(size)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, 18, 10, GREEN_DARK),
        "Heading 2": (13, 14, 7, GREEN),
        "Heading 3": (11.5, 10, 5, INK),
    }
    for name, (size, before, after, color) in heading_specs.items():
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)

    # User preference: no running header or footer on generated DOCX reports.
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        section.header.paragraphs[0].text = ""
        section.footer.paragraphs[0].text = ""

    core = doc.core_properties
    core.title = "PROJECT MIRROR Steam正式发售级UI交互改进与最终设计报告"
    core.subject = "UI/UX审计、交互改进与最终设计规范"
    core.author = "PROJECT MIRROR Design Review"
    core.keywords = "PROJECT MIRROR, UI, UX, Steam, Steam Deck, Godot"


def add_cover(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(52)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_run = kicker.add_run("PROJECT MIRROR · PRODUCT DESIGN AUDIT")
    font_run(kicker_run, size=10, bold=True, color=GREEN)
    kicker.paragraph_format.space_after = Pt(16)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Steam 正式发售级\nUI 交互改进与最终设计报告")
    font_run(title_run, size=28, bold=True, color=INK)
    title.paragraph_format.space_after = Pt(14)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("基于当前 Godot Demo 实机流程、审核后正式设计与绿色激光 UI 规范")
    font_run(subtitle_run, size=12.5, color=GRAY)
    subtitle.paragraph_format.space_after = Pt(28)

    line = doc.add_paragraph()
    set_paragraph_border_bottom(line)
    line.paragraph_format.space_after = Pt(26)

    add_figure(
        doc,
        RUNTIME / "03-bash-gameplay.png",
        "当前构建实机画面：Bash 玩家回合（本轮重新渲染并检查）",
        width=6.25,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"版本：设计评审稿 1.0  |  日期：{date.today().isoformat()}  |  目标平台：Windows PC / Steam Deck"
    )
    font_run(meta_run, size=9, color=GRAY)
    doc.add_page_break()


def build_report() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "执行结论", 1)
    add_callout(
        doc,
        "总体判断",
        "当前 Demo 已经是功能完整、可验证的垂直切片，但还不是可直接面向 Steam 玩家销售的表现层。"
        "最优策略不是推翻绿色激光与浅色玻璃体系，而是把工程验证界面升级为有角色、有节奏、有反馈、可完整设置的商业游戏 UI。",
    )
    add_body(
        doc,
        "现有实现最强的部分是流程约束：TitleScreen、背景/教学、Bash 两轮、Limit Bash、结果、Summary 与稳定存档之间的关系清晰；"
        "三选一、确认、锁定、揭示、行动日志和恢复也都有真实实现。最弱的部分是玩家感知层：大面积静态空白、Tutor 字母占位、"
        "同一玻璃卡片贯穿所有场景、结果与对话缺少表演、设置/暂停/输入提示缺位，让完成度更接近内部 QA 工具。"
    )
    add_table(
        doc,
        ["维度", "当前状态", "Steam 发售目标", "优先级"],
        [
            ("流程完整性", "8.5/10：主流程与存档扎实", "保留现有状态机与稳定检查点", "保持"),
            ("视觉层级", "5/10：结构清楚但空白过多", "中心玩法、Tutor、状态与主 CTA 一眼可辨", "P0"),
            ("交互反馈", "5.5/10：首页有基础动效，局内偏静态", "所有状态都有视觉/音频/手柄反馈", "P0"),
            ("叙事演出", "3/10：T 字母占位，文本卡片静态", "角色反应、镜头、对话节奏与章节过渡完整", "P0"),
            ("设置与暂停", "1/10：Settings 隐藏，无正式暂停层", "主菜单与暂停菜单均可访问完整设置", "P0"),
            ("控制器 / Deck", "2/10：无完整输入映射与动态字形", "全流程手柄可用、1280×800 易读、字形随设备切换", "P0"),
            ("可访问性", "2.5/10：颜色带文字，但缺少可调选项", "文本、对比度、动效、闪烁、字幕、输入均可调整", "P0"),
            ("本地化", "2/10：主要为硬编码英文", "英/简中/日文外置资源与 30% 扩展空间", "P1"),
        ],
        [1800, 2500, 3560, 1500],
    )

    add_heading(doc, "1. 审计范围与证据", 1)
    add_body(
        doc,
        "审计对象包括 demo 当前 C#/Godot 场景与主题、审核后正式设计中的 Demo设计.docx / DemoUI.pptx、绿色激光版 24 状态规范，"
        "以及本轮重新运行的实际 1920×1080 流程截图。旧 QA 图仅作为回滚备份，不作为本轮审计结论的唯一证据。"
    )
    add_table(
        doc,
        ["步骤", "玩家看到的状态", "健康度", "本轮证据"],
        [
            ("1", "无存档 TitleScreen", "良好：入口明确；弱点是设置缺位、面板过大", "01-title.png"),
            ("2", "Chapter 0 整页章节卡", "一般：清晰但缺少推进提示与场景演出", "02-chapter-0.png"),
            ("3", "Tutor 打字机对话", "一般：机制可用；首帧单字符和巨量空白影响观感", "02b / 02c"),
            ("4", "Bash 玩家回合", "良好：规则、剩余、选择、日志并列；信息量偏工程化", "03-bash-gameplay.png"),
            ("5", "Bash 选中态", "良好：选中与确认清楚；缺少音效/手柄震动/动作预演", "03a-bash-selected.png"),
            ("6", "Bash / Limit 结算", "一般：数据完整；日志夺取注意力，结果缺少角色表演", "03b / 03c"),
            ("7", "稳定存档恢复", "良好：恢复提示、进度与时间明确", "04-bash-restored.png"),
            ("8", "Summary", "一般：统计正确；结束仪式感、下一步和分享/重玩入口不足", "04-summary.png"),
        ],
        [850, 2900, 2520, 3090],
    )
    add_two_figures(
        doc,
        RUNTIME / "01-title.png",
        APPROVED_SLIDES / "slide-2.png",
        "左：当前无存档首页；右：审核后正式设计首页。结构已高度一致，但正式构建仍缺设置、版本/隐私信息与更强的视觉焦点。",
    )
    add_two_figures(
        doc,
        RUNTIME / "04-summary.png",
        APPROVED_SLIDES / "slide-11.png",
        "左：当前 Summary；右：审核后设计。当前统计更真实，但仍缺章节进度、角色反应与完成后的明确下一步。",
    )

    add_heading(doc, "2. 当前 Demo 已实现的 UI 功能", 1)
    add_heading(doc, "2.1 进入、存档与路由", 2)
    for item in (
        "TitleScreen 启动即显示；New Game、有效存档 Continue、Quit 均已接入实际路由。",
        "存在未完成存档时，New Game 会先显示覆盖确认；Continue 显示存档短 ID、阶段与更新时间。",
        "Continue 恢复前锁定按钮，损坏存档能够回退/警告；完成存档后隐藏 Continue。",
        "首页按钮支持鼠标悬浮、键盘焦点和 0.98 压缩回弹；ReducedMotion 属性存在，但尚未形成全局设置链路。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2.2 对话、章节与教学", 2)
    for item in (
        "Chapter 0 / Chapter 1 使用独立整页覆盖层，纯说明页隐藏三选一与确认按钮。",
        "Tutor 对话采用每秒 10 字符打字机；第一次左键完成当前行，第二次推进；页面完成后写入稳定检查点。",
        "对话页顶部显示阶段、存档页码与实际游玩时间；Summary 复用独立对话页面。",
        "当前推进仅明确支持左键；UIAccept、手柄 A、键盘 Space/Enter 与长按跳过没有形成统一输入动作。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2.3 Bash 与 Limit Bash 局内交互", 2)
    add_table(
        doc,
        ["状态", "已实现表现", "当前限制"],
        [
            ("Idle", "三按钮固定等宽；非法选项禁用；确认按钮提示先选择", "颜色与形状清楚，但缺动态输入字形"),
            ("Selected", "提高填充、5 px 描边、对勾、状态文字、Confirm 启用", "缺选择音、触觉、预览动画"),
            ("Locked", "按钮失焦/不可用；Limit 显示等待双方", "锁定只靠静态灰色和文字，缺阶段性演出"),
            ("Reveal", "ShowLimitReveal 同时显示 Player/Tutor 选择并保留回合日志", "缺可见的双侧揭示镜头与数量变化动画"),
            ("Result", "锁定输入；显示胜负、累计数据、最终选择和 Continue", "角色无表情，结果卡与日志竞争注意力"),
            ("Restore", "重建规则状态、日志、时间和玩家可输入状态", "恢复提示仅文本，没有短暂重建动效"),
        ],
        [1250, 4050, 4060],
    )
    add_figure(
        doc,
        RUNTIME / "03a-bash-selected.png",
        "当前选中态：固定三等分、对勾、深描边与独立确认按钮已经满足基础可用性。",
    )

    add_heading(doc, "3. 距离正式发售的主要问题", 1)
    add_heading(doc, "3.1 P0：会直接影响首发评价的问题", 2)
    p0_items = [
        ("Tutor 缺少正式角色表现", "字母 T 仍是开发占位，玩家无法形成情感连接；对话、胜负、锁定与恢复没有角色反应。"),
        ("信息架构偏调试面板", "左右栏长期同时显示完整状态和 SYSTEM 日志，中心玩法被压缩，玩家先看到系统而不是游戏。"),
        ("没有正式暂停与设置", "Save & Back 常驻底部破坏沉浸；Settings 节点隐藏，音频、显示、语言、输入、无障碍均不可调。"),
        ("输入体系未闭环", "章节与对话依赖左键；没有动态键鼠/手柄字形、统一焦点规则、重映射和 Steam Input 行为。"),
        ("场景转换是硬切", "章节卡、对话、游戏、结算和背景变化缺少连续镜头，用户感知为多张静态页面依次出现。"),
        ("对话节奏过慢且缺提示", "默认 10 cps 会让长句等待明显；首帧只出现 1 个字符，页面没有‘点击完成/继续’的反馈。"),
        ("Steam Deck / 16:10 未专项设计", "当前以 1920×1080 为唯一视觉目标，右侧日志与三栏布局在 1280×800 需要折叠策略。"),
        ("本地化与文本扩展不足", "英文硬编码与 JSON 混用；中文、日文、德文扩展会改变行数和按钮宽度。"),
    ]
    for title, detail in p0_items:
        add_body(doc, f"{title}：{detail}", bold_lead=f"{title}：")

    add_heading(doc, "3.2 P1/P2：从可售到出色", 2)
    add_body(
        doc,
        "P1 应补齐角色动画、局内拿取表现、Reveal 编排、结果镜头、统一音效、手柄轻震、局内日志抽屉、"
        "设置实时预览和本地化。P2 再加入更细的玻璃折射、环境光、叙事化背景切换、成就/Steam Overlay 提示与更丰富的 Tutor 表情。"
    )

    add_heading(doc, "4. 最终设计方向", 1)
    add_callout(
        doc,
        "设计主张",
        "保留‘冷白实验室 + 浅色玻璃 + 绿色激光’作为品牌骨架；减少同时可见的工程数据，把注意力顺序固定为："
        "玩家目标 → 当前状态 → 可行动作 → Tutor 反馈 → 可选详细记录。",
    )
    add_table(
        doc,
        ["令牌", "最终值", "用途"],
        [
            ("主色", "Laser Green #26D89A / Deep Green #0B7055", "焦点、确认、稳定状态、结构线"),
            ("表面", "Glass #F4FBF8E6 / Elevated #FFFFFFF2", "主玻璃面板与高层弹窗"),
            ("正文", "Ink #123D35 / Secondary #5A6B67", "确保浅色玻璃上的持续可读性"),
            ("风险", "Amber #E6B94A / Coral #E56375", "等待/警告/错误；必须配文字或图标"),
            ("锁定", "Neutral #95A4A0", "不可交互，不改变原始按钮文案"),
            ("圆角", "卡片 24；按钮 14；模态 28", "保持现有柔和玻璃语言"),
            ("字体", "Inter Variable + Noto Sans CJK", "标题紧凑、正文多语一致；允许 100/125/150%"),
            ("栅格", "1920×1080：12 列、80 边距、24 沟槽", "统一对齐；16:10 使用 48 px 安全边距"),
        ],
        [1500, 3400, 4460],
    )

    add_heading(doc, "5. 最终 UI 界面排布", 1)
    add_heading(doc, "5.1 全局架构", 2)
    add_table(
        doc,
        ["区域", "1920×1080 桌面", "1280×800 Steam Deck", "交互规则"],
        [
            ("顶部状态条", "72–80 px；阶段/回合在左，计时/保存/菜单在右", "64 px；只保留阶段与菜单", "不放胜负详情；状态变化以文字+图标表达"),
            ("左侧规则区", "280–320 px 常驻，可折叠", "默认收起为 Rules 按钮", "只显示当前规则与 1 条关键提示"),
            ("中央玩法区", "720–900 px，包含实体化剩余量、Tutor 反馈、选择", "占主要宽度；优先保证选择与确认", "主 CTA 始终处于视觉轴线"),
            ("右侧记录区", "320–360 px；默认只显示最近 2–3 条", "抽屉式 Log，按 Y/Tab 打开", "完整审计日志可访问但不抢视线"),
            ("底部动作区", "三等分选择 + 独立 Confirm；高度 64–72 px", "高度 56–64 px；A 确认、B 返回", "命中区不随文字扩展；提示字形动态切换"),
        ],
        [1400, 2800, 2460, 2700],
    )

    add_heading(doc, "5.2 各界面最终设计", 2)
    screen_specs = [
        ("TitleScreen", "左侧 520–600 px 菜单玻璃；右侧保留实验室主体与 Tutor 模糊轮廓。New Game / Continue 为主操作，Settings / Accessibility / Quit 放次级行。存档卡显示阶段、游玩时长、最后更新时间，不暴露原始 ID。"),
        ("Chapter Transition", "不再是单一空白卡。背景轻度景深推进，章节号从 8% 透明度滑入；副标题与一句目标同步出现；底部显示 Continue / Hold to Skip。预载完成后才允许进入。"),
        ("Dialogue Only", "Tutor 正式角色占左侧 32–36%；对话卡位于右下并限制为 2–4 行。顶部仅保留阶段与存档状态；右下持续显示 Advance 字形、Auto、Skip 和速度。"),
        ("Bash Gameplay", "中央把 Remaining 从纯数字升级为可数的粒子/刻度环；玩家或 Tutor 拿取时实体减少。规则栏仅保留当前胜负条件；日志折叠。选择与 Confirm 位于底部主轴。"),
        ("Limit Bash Lock / Reveal", "锁定后玩家牌面移到左侧保险框，Tutor 牌面以扫描遮罩出现在右侧；双方 Ready 后进行 0.8 秒同场揭示。Reveal 完成前不更新 Remaining。"),
        ("Round Result", "中心结果牌 + Tutor 表情/姿态 + 三个关键数字；完整回合日志在二级抽屉。Continue 为唯一主 CTA，Replay/View Log 为次级。"),
        ("Summary", "用阶段时间线展示 Background → Bash R1 → Bash R2 → Limit → Complete；保留胜负/局数/回合数和完成原因。返回标题、重新开始、查看统计分层。"),
        ("Pause / Settings", "Esc/Menu 打开暂停层：Resume、Settings、Controls、Save & Return、Quit to Desktop。设置含显示、音频、控制、无障碍、语言、隐私，并提供实时预览。"),
        ("Modal / Error", "覆盖存档、损坏存档、写入失败与联网错误使用统一 560–680 px 模态；标题说明结果，正文说明恢复动作，焦点默认落在安全选项。"),
    ]
    for title, detail in screen_specs:
        add_body(doc, f"{title}：{detail}", bold_lead=f"{title}：")

    add_figure(
        doc,
        GREEN_SPEC_SLIDES / "slide-3.png",
        "绿色激光版共享栅格：最终设计继续沿用三栏骨架，但将右侧记录改为渐进披露，把中央玩法恢复为第一焦点。",
    )

    add_heading(doc, "6. 最终交互方式", 1)
    add_table(
        doc,
        ["动作", "键鼠", "控制器 / Deck", "行为"],
        [
            ("移动焦点", "方向键 / WASD", "D-pad / 左摇杆", "焦点按视觉顺序移动，不进入锁定/隐藏控件"),
            ("选择", "左键 / Space", "A / Cross", "仅改变 Selected；可反复改选"),
            ("确认", "Enter / 独立 Confirm", "A（焦点在 Confirm）", "提交后立即锁定；忽略重复输入"),
            ("取消/返回", "Esc / 右键", "B / Circle", "对话返回暂停；游戏中打开 Pause，不直接丢失进度"),
            ("对话推进", "左键 / Space / Enter", "A", "第一次补完文字，第二次推进；长按 0.45 s 开启快速推进"),
            ("日志", "Tab / L", "Y / Triangle", "打开可滚动抽屉；关闭后焦点回到原控件"),
            ("无障碍快捷", "F10", "长按 Menu", "打开快速可访问设置：字号、对比度、动效、字幕"),
        ],
        [1300, 2250, 2250, 3560],
    )
    add_heading(doc, "6.1 通用控件状态", 2)
    add_table(
        doc,
        ["状态", "视觉", "声音/触觉", "时序与恢复"],
        [
            ("Idle", "浅玻璃、低亮边框", "无", "静态；保持稳定命中区"),
            ("Hover / Focus", "亮度 +8%，3 px 焦点环，上浮 2 px", "轻 UI tick；手柄不震动", "160–220 ms ease-out"),
            ("Selected", "风险色浅填充 + 深描边 + 对勾/标签", "select 音；10–15 ms 轻震", "120–160 ms；确认前可逆"),
            ("Confirmed / Locked", "灰化未选项，已选项进入保险框", "lock 音；25 ms 脉冲", "90–120 ms；输入路由锁定"),
            ("Processing", "扫描线/三点状态，保留原值", "低频循环，1.5 s 后才出现", "不可无限；超时进入可恢复错误"),
            ("Reveal", "双方牌面同帧翻开，Remaining 延后更新", "双重揭示音 + 40 ms 震动", "650–900 ms；Reduced Motion 180 ms 淡变"),
            ("Result", "胜负大字、角色反应、关键数字", "不同胜/负/平短音型", "900–1200 ms；主 CTA 在 600 ms 后可用"),
            ("Error", "红色图标 + 原因 + 恢复动作", "单次低音，不循环", "焦点留在重试/安全返回"),
        ],
        [1250, 3350, 2300, 2460],
    )

    add_heading(doc, "7. 对话出现与 UI 动效", 1)
    add_heading(doc, "7.1 对话序列", 2)
    dialogue_sequence = [
        "0–120 ms：对话卡从 96% 缩放和 8 px 下移进入稳定位置；角色保持当前表情。",
        "120–260 ms：Speaker、情绪标签与推进字形淡入；正文仍隐藏，避免单字符孤立出现。",
        "260 ms 起：正文默认 34 cps；标点停顿 80–140 ms；允许设置为 20/34/60/Instant。",
        "玩家第一次推进：在 80 ms 内完成本行并播放极轻确认音；第二次推进下一行。",
        "最后一行完成：右下 Advance 字形进行 1.6 s 呼吸，不让整张卡片跳动；Auto 模式按阅读长度停留。",
        "Reduced Motion：取消位移/缩放，只保留 120–180 ms 透明度变化；Instant Text 直接显示全文。",
    ]
    for item in dialogue_sequence:
        add_number(doc, item)

    add_heading(doc, "7.2 Tutor 表现", 2)
    for item in (
        "正式资产应替代字母 T：建议使用半透明液体球/合成生物头像，保持与现有设计稿的圆形构图一致。",
        "Idle：2.8–4.2 s 随机呼吸、眼动与弱扫描纹；不持续闪烁。",
        "Thinking：眼神转向棋盘，绿色弧线沿轮廓运行一次；超过 1.5 s 才显示文字状态。",
        "Lock：角色与牌面同步短促凝固；Reveal：转向玩家选择；Result：胜、负、平使用不同但克制的反应。",
        "对话情绪不改变正文位置；角色动画可单独关闭，字幕与关键状态始终可用。",
    ):
        add_bullet(doc, item)

    add_figure(
        doc,
        RUNTIME / "02c-background-02.png",
        "当前对话页的主要问题：Tutor 仍为占位符，文字卡面积过大且没有持续的推进提示。",
    )

    add_heading(doc, "8. 场景切换效果", 1)
    add_table(
        doc,
        ["切换", "标准动效", "Reduced Motion", "输入门槛"],
        [
            ("Title → Chapter", "650 ms 玻璃门合拢 + 绿色扫描线；背景景深切换", "180 ms 交叉淡变", "新场景稳定、资源加载完成后解锁"),
            ("Chapter → Dialogue", "章节标题向上缩为阶段标签，Tutor 从景深中显现", "标题淡出/对话淡入", "允许 A/Enter 跳过章节卡"),
            ("Dialogue → Gameplay", "对话卡下沉，中央 Remaining 从背景粒子聚合", "150 ms 淡变", "选择控件出现后默认焦点落在首个合法动作"),
            ("Bash → Limit", "三按钮短暂熄灭；交替箭头拆分为双侧保险框", "状态文本直接替换", "玩家必须确认读完规则变化"),
            ("Lock → Reveal", "双侧 3-2-1 扫描；同帧翻开选择；再扣减 Remaining", "同帧文本与数值淡变", "Reveal 结束前忽略全部选择输入"),
            ("Result → Next", "结果牌收进阶段时间线，背景色温轻微变化", "结果淡出/新局淡入", "存档写入成功或可恢复失败已提示"),
            ("Summary → Title", "完成标记写入后，时间线化为 TitleScreen 的存档状态", "180 ms 淡变", "完成存档必须先关闭 Continue"),
        ],
        [1700, 3400, 2050, 2210],
    )

    add_heading(doc, "9. UI 特效、音频与触觉", 1)
    add_heading(doc, "9.1 UI 特效", 2)
    for item in (
        "玻璃：背景模糊 12–18 px、面板透明 86–94%、边缘仅 1–2 px；不在正文后持续移动高亮。",
        "绿色激光：仅用于焦点、确认、锁定和揭示路径；单次运动 280–420 ms，避免常驻霓虹造成疲劳。",
        "剩余量：用 12–30 个可数刻度/粒子表达，数字作为辅助；拿取时从右向左熄灭并保留剩余数字。",
        "风险按钮：绿/黄/红继续保留，但增加 1/2/3、低/中/高或形状纹理，避免颜色成为唯一信息。",
        "错误：禁止整屏红闪；使用局部红边、静态图标和可恢复说明。",
        "性能：玻璃与粒子分为 High / Balanced / Low；Steam Deck 默认 Balanced，保证 UI 动效稳定。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "9.2 音频与触觉", 2)
    add_table(
        doc,
        ["事件", "音频", "触觉", "可调性"],
        [
            ("Hover/Focus", "短促玻璃 tick，避免每帧重复", "无", "UI 音量"),
            ("Select", "三选项使用同一音高族", "10–15 ms", "UI 音量 / 触觉强度"),
            ("Confirm/Lock", "低频锁定 + 轻金属闭合", "20–30 ms", "可关闭"),
            ("Reveal", "左右声像先分离再汇合", "35–45 ms", "可关闭立体声增强"),
            ("Win/Lose/Draw", "三种独立 0.8–1.2 s 音型", "40–80 ms，强度不同", "结果音量 / 触觉强度"),
            ("Dialogue", "每 2–4 字符一次极轻合成语音粒子", "无", "语音/文本音独立关闭"),
        ],
        [1700, 3800, 1500, 2360],
    )

    add_heading(doc, "10. 设置、可访问性、本地化与 Steam Deck", 1)
    add_heading(doc, "10.1 正式设置结构", 2)
    add_table(
        doc,
        ["分类", "必须包含"],
        [
            ("Display", "分辨率、窗口/全屏、VSync、UI 缩放、背景模糊、特效质量、亮度/对比度"),
            ("Controls", "键鼠/手柄自动识别、重映射、震动、字形类型、焦点高亮、长按时间"),
            ("Audio", "Master、BGM、Voice/Text、UI、Result；独立静音"),
            ("Accessibility", "字号 100/125/150%、高对比度、色彩替代、Reduced Motion、减少闪烁、Instant Text、字幕背景"),
            ("Language", "English、简体中文、日本語；即时预览并在安全点应用"),
            ("Privacy / Data", "本地存档说明、清除记录、联网评价状态；DEMO 保持离线边界"),
        ],
        [1850, 7510],
    )
    add_heading(doc, "10.2 Steam/Deck 目标", 2)
    for item in (
        "支持 1280×800（优先）与 1280×720；关键正文按 12 px 以上目标设计，不把 9 px 作为常态。",
        "默认控制器配置必须覆盖全部游戏功能；键鼠与手柄可以随时混用；提示字形随当前设备切换。",
        "所有焦点都能在手柄上到达并看见；打开日志、设置、确认覆盖和恢复错误均不得要求鼠标。",
        "正文和按钮预留至少 30% 本地化扩展；UI 缩放不改变三选一顺序和确认按钮位置。",
        "Steam 商店可访问性标签只能在功能真实完成并通过测试后勾选，不能把设计计划当成已支持。",
    ):
        add_bullet(doc, item)
    add_callout(
        doc,
        "外部标准边界",
        "Steamworks 当前明确要求/建议包括完整控制器访问、正确输入字形、Steam Deck 1280×800 适配、"
        "文字可读性和可调字号/对比度。本文把这些作为 UI 设计验收目标，不宣称当前构建已经通过 Steam Deck Verified 或完整无障碍认证。",
        fill=YELLOW_PALE,
        color=YELLOW,
    )

    add_heading(doc, "11. 实施路线图", 1)
    add_table(
        doc,
        ["阶段", "交付内容", "完成定义"],
        [
            ("P0-A 信息架构", "日志渐进披露、规则栏压缩、暂停/设置、Title 存档卡", "核心流程在 1920×1080 / 1280×800 均无拥挤"),
            ("P0-B 输入与访问", "InputMap、手柄焦点、动态字形、字号/动效/对比度设置", "仅用手柄完成新游戏到 Summary"),
            ("P0-C 叙事表现", "正式 Tutor 资产、对话卡、推进提示、章节/场景过渡", "不再出现字母 T 占位与硬切页面"),
            ("P1 动效与反馈", "选择/锁定/Reveal/Result 动画、音效、触觉", "每个交互状态都有可见且可关闭的反馈"),
            ("P1 本地化", "英/简中/日文资源化、字体、文本扩展测试", "三语言完整流程无截断、错位与硬编码"),
            ("P2 Steam 集成", "Steam Input、Deck 设备测试、Overlay/成就提示", "通过内部 Deck 清单，再提交官方兼容性评审"),
        ],
        [1400, 4400, 3560],
    )

    add_heading(doc, "12. 最终验收标准", 1)
    acceptance = [
        "从启动到 Summary 的所有主操作可仅用 Xbox/PlayStation/Steam Deck 控制器完成。",
        "在 1920×1080、2560×1440、1280×800、1280×720 下无裁切、重叠、焦点丢失或必须滚动的主 CTA。",
        "正文默认可读，Deck 目标最小字符高度 12 px；字号 150% 时仍保持选择/确认可达。",
        "颜色不作为唯一状态：Selected、Locked、Error 均同时有文字、形状或图标。",
        "Reduced Motion 能关闭位移、缩放、震动和闪烁；Instant Text 能取消打字机。",
        "对话第一次推进补完文本，第二次推进；键鼠与手柄行为一致；长按跳过可撤销。",
        "Limit Bash 的锁定、双方 Ready、同场揭示、Remaining 更新具有明确顺序，且 Reveal 前不能重复提交。",
        "完整日志可访问，但默认界面只显示最近事件；结果页的第一焦点是结果与 Continue。",
        "暂停/设置可从 Title 和局内进入；改变音频、字幕、字号、动效后立即预览并持久化。",
        "英/简中/日文完整通关无硬编码残留；输入提示随设备切换且不显示错误字形。",
        "存档覆盖、恢复失败、写入失败和完成存档都使用统一可恢复模态并有安全默认焦点。",
        "每个关键页面完成截图回归、键盘/手柄走查、色彩/对比度审查和性能检查。",
    ]
    for item in acceptance:
        add_bullet(doc, item, compact=True)

    add_heading(doc, "13. 证据限制", 1)
    add_body(
        doc,
        "本报告对 UX、视觉与可访问性风险的判断来自本轮实际截图、代码/场景结构和现有设计文件。"
        "截图能够证明布局、文本、状态和部分焦点反馈，但不能单独证明完整键盘/手柄顺序、颜色对比数值、屏幕阅读器兼容性、"
        "音频/触觉质量或 Steam Deck 性能。上述项目必须在实现后使用真实设备和自动化/人工测试验证。"
    )

    add_heading(doc, "参考资料", 1)
    references = [
        ("PROJECT MIRROR 审核后正式设计：Demo设计.docx", r"Documents\审核后正式设计\Demo设计.docx"),
        ("PROJECT MIRROR 审核后 UI：DemoUI.pptx", r"Documents\审核后正式设计\DemoUI.pptx"),
        ("绿色激光 UI 设计规范 24 状态稿", r"Documents\AI生成的UI\PROJECT_MIRROR_UI设计规范_绿色激光版_重设计.pptx"),
        ("Steamworks：Steam Deck and Steam Machine Compatibility Review", "https://partner.steamgames.com/doc/steamhardware/compat?language=english"),
        ("Steamworks：Getting Started for Steam Input Developers", "https://partner.steamgames.com/doc/features/steam_controller/getting_started_for_devs?language=english"),
        ("Steamworks：Accessibility Features", "https://partner.steamgames.com/doc/accessibility_features?language=english"),
        ("Steamworks：Localization and Languages", "https://partner.steamgames.com/doc/store/localization?language=english"),
    ]
    for label, target in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        bullet = paragraph.add_run("• ")
        font_run(bullet, color=GREEN)
        if target.startswith("https://"):
            add_hyperlink(paragraph, label, target)
        else:
            run = paragraph.add_run(f"{label} — {target}")
            font_run(run, size=9.0, color=INK)

    add_callout(
        doc,
        "最终设计判断",
        "正式版不需要抛弃现有 Demo。保留规则、存档与绿色激光骨架，优先重做信息层级、输入设置、Tutor 表现和场景时序，"
        "即可把当前‘可证明功能’的垂直切片升级为‘玩家愿意购买并持续体验’的产品界面。",
        trailing_space=False,
    )
    return doc


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    report = build_report()
    report.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
